# -*- coding: utf-8 -*-
"""
Sinh file DOCX: FORM HOÀN THIỆN & NÂNG CẤP WEBSITE Y VIỆN TOPLINK
Một biểu mẫu điền tay, bao trùm mọi mặt (trừ database).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Bảng màu thương hiệu ----------
CRIMSON   = RGBColor(0x95, 0x13, 0x1f)
CRIMSON_D = RGBColor(0x5c, 0x0a, 0x11)
GOLD      = RGBColor(0xb0, 0x86, 0x1e)
WOOD      = RGBColor(0x50, 0x2c, 0x1e)
INK       = RGBColor(0x31, 0x31, 0x31)
INK_SOFT  = RGBColor(0x5c, 0x53, 0x4c)
WHITE     = RGBColor(0xff, 0xff, 0xff)
FILL_GREY = "F2EFE9"   # ô điền
FILL_HEAD = "95131F"   # nền tiêu đề bảng
FILL_NOTE = "FBF3E6"   # nền ghi chú
FILL_SUB  = "EDE7D3"   # nền tiểu mục

BODY_FONT = "Be Vietnam Pro"
HEAD_FONT = "Times New Roman"  # thay cho Cormorant/Playfair (an toàn khi mở máy khác)

doc = Document()

# ---------- style mặc định ----------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
pf = normal.paragraph_format
pf.space_after = Pt(4)
pf.line_spacing = 1.12

for sec in doc.sections:
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

# ---------- helpers ----------
def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def _no_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "D8CFC0")
        borders.append(e)
    tblPr.append(borders)

def run(p, text, bold=False, italic=False, size=10.5, color=INK, font=BODY_FONT, caps=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if caps:
        rPr = r._element.get_or_add_rPr()
        c = OxmlElement("w:caps"); c.set(qn("w:val"), "true"); rPr.append(c)
    return r

def spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p

# ---- Tiêu đề Phần lớn (A, B, C...) ----
def part_heading(letter, title):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    _shade(c, FILL_HEAD)
    _set_cell_margins(c, 120, 120, 160, 160)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, f"PHẦN {letter}", bold=True, size=11, color=RGBColor(0xf3,0xd2,0x7a), font=BODY_FONT, caps=True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run(p2, title, bold=True, size=15, color=WHITE, font=HEAD_FONT)
    spacer(6)

# ---- Tiêu đề mục con (số) ----
def sub_heading(num, title, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "C8A02E")
    pbdr.append(bottom); pPr.append(pbdr)
    run(p, f"{num}  ", bold=True, size=13, color=GOLD, font=HEAD_FONT)
    run(p, title, bold=True, size=13, color=CRIMSON, font=HEAD_FONT)
    if desc:
        d = doc.add_paragraph()
        d.paragraph_format.space_after = Pt(4)
        run(d, desc, italic=True, size=9.5, color=INK_SOFT)

def note(text, label="Gợi ý"):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    _shade(c, FILL_NOTE)
    _set_cell_margins(c, 70, 70, 130, 130)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, f"{label}: ", bold=True, size=9.5, color=GOLD)
    run(p, text, italic=True, size=9.5, color=INK_SOFT)
    spacer(4)

def para(text, bold=False, size=10.5, color=INK, italic=False):
    p = doc.add_paragraph()
    run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

# ---- Bảng điền: list[(field, hint)] -> 2 cột (Mục | điền) ----
def fill_table(rows, col1="Mục cần điền", col2="Nội dung điền", w1=4.6, w2=10.4):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_border(t)
    t.autofit = False
    hdr = t.rows[0].cells
    for i, txt in enumerate((col1, col2)):
        _shade(hdr[i], FILL_HEAD)
        _set_cell_margins(hdr[i])
        pp = hdr[i].paragraphs[0]; pp.paragraph_format.space_after = Pt(0)
        run(pp, txt, bold=True, size=10, color=WHITE)
    for field, hint in rows:
        cells = t.add_row().cells
        _set_cell_margins(cells[0]); _set_cell_margins(cells[1])
        _shade(cells[1], FILL_GREY)
        p0 = cells[0].paragraphs[0]; p0.paragraph_format.space_after = Pt(0)
        run(p0, field, bold=True, size=10, color=INK)
        if hint:
            ph = cells[0].add_paragraph(); ph.paragraph_format.space_after = Pt(0)
            run(ph, hint, italic=True, size=8.5, color=INK_SOFT)
        cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    # đặt độ rộng cột
    for row in t.rows:
        row.cells[0].width = Cm(w1)
        row.cells[1].width = Cm(w2)
    spacer(6)
    return t

# ---- Bảng nhiều cột để liệt kê (vd danh sách dịch vụ) ----
def grid_table(headers, n_blank_rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_border(t)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], FILL_HEAD); _set_cell_margins(hdr[i])
        pp = hdr[i].paragraphs[0]; pp.paragraph_format.space_after = Pt(0)
        run(pp, h, bold=True, size=9.5, color=WHITE)
    for _ in range(n_blank_rows):
        cells = t.add_row().cells
        for j, c in enumerate(cells):
            _set_cell_margins(c)
            _shade(c, FILL_GREY)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    spacer(6)
    return t

# ---- Câu hỏi lựa chọn (checkbox) ----
def choice(question, options, multi=False, desc=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, "▸ ", bold=True, color=CRIMSON)
    run(p, question, bold=True, size=10.5, color=INK)
    if desc:
        run(p, f"  ({desc})", italic=True, size=9, color=INK_SOFT)
    note_txt = "chọn nhiều" if multi else "chọn một"
    for opt in options:
        po = doc.add_paragraph(); po.paragraph_format.space_after = Pt(1)
        po.paragraph_format.left_indent = Cm(0.6)
        run(po, "☐  ", size=11, color=GOLD)
        run(po, opt, size=10, color=INK)
    pe = doc.add_paragraph(); pe.paragraph_format.space_after = Pt(1)
    pe.paragraph_format.left_indent = Cm(0.6)
    run(pe, "☐  Khác: ", size=11, color=GOLD)
    run(pe, "______________________________________________", size=10, color=INK_SOFT)
    pn = doc.add_paragraph(); pn.paragraph_format.space_after = Pt(6)
    pn.paragraph_format.left_indent = Cm(0.6)
    run(pn, f"({note_txt})", italic=True, size=8, color=INK_SOFT)

# ---- Khối điền tự do nhiều dòng ----
def open_block(label, lines=3, hint=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, "▸ ", bold=True, color=CRIMSON)
    run(p, label, bold=True, size=10.5, color=INK)
    if hint:
        run(p, f"  — {hint}", italic=True, size=9, color=INK_SOFT)
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    _shade(c, FILL_GREY); _set_cell_margins(c, 90, 90, 130, 130)
    cp = c.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
    for _ in range(lines - 1):
        c.add_paragraph().paragraph_format.space_after = Pt(0)
    spacer(6)

# ============================================================
#  TRANG BÌA
# ============================================================
for _ in range(2):
    spacer(10)
pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(pc, "Y VIỆN TOPLINK", bold=True, size=12, color=GOLD, caps=True)
pc.paragraph_format.space_after = Pt(2)

pt = doc.add_paragraph(); pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(pt, "FORM HOÀN THIỆN & NÂNG CẤP WEBSITE", bold=True, size=26, color=CRIMSON, font=HEAD_FONT)
pt.paragraph_format.space_after = Pt(4)

ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(ps, "Biểu mẫu thu thập toàn bộ thông tin & quyết định thiết kế — từ chi tiết nhỏ đến lớn",
    italic=True, size=12, color=INK_SOFT, font=HEAD_FONT)
ps.paragraph_format.space_after = Pt(2)
ps2 = doc.add_paragraph(); ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(ps2, "Phong cách thương hiệu: Đông Y Cung Đình Hiện Đại  ·  (không bao gồm phần database)",
    italic=True, size=10.5, color=GOLD)

spacer(14)

# Khung hướng dẫn
t = doc.add_table(rows=1, cols=1)
c = t.cell(0, 0); _shade(c, FILL_NOTE); _set_cell_margins(c, 150, 150, 180, 180)
hp = c.paragraphs[0]
run(hp, "CÁCH DÙNG BIỂU MẪU NÀY", bold=True, size=11, color=CRIMSON); hp.paragraph_format.space_after = Pt(4)
for line in [
    "• Điền trực tiếp vào ô màu xám bên cạnh/bên dưới mỗi mục. Tích ☐ thành ☒ với câu hỏi lựa chọn.",
    "• Mục nào chưa quyết được, ghi “để Toplink đề xuất” — đội thiết kế sẽ chốt theo bộ nhận diện.",
    "• Phần A–C là nền tảng (thương hiệu, nhận diện, liên hệ). Phần D–E quyết định cấu trúc & từng section.",
    "• Phần F–J là nội dung chi tiết, hình ảnh, chuyển đổi, SEO và giọng văn.",
    "• Mọi mã màu/HEX, font, bo góc, viền, hiệu ứng đang dùng đều có sẵn để bạn xác nhận hoặc đổi.",
]:
    lp = c.add_paragraph(); lp.paragraph_format.space_after = Pt(2)
    run(lp, line, size=10, color=INK)
mp = c.add_paragraph(); mp.paragraph_format.space_after = Pt(0)
run(mp, "Ngày điền: ______/______/__________            Người điền: ______________________________",
    size=10, color=INK_SOFT)

doc.add_page_break()

# ============================================================
#  MỤC LỤC (thủ công)
# ============================================================
para("MỤC LỤC CÁC PHẦN", bold=True, size=14, color=CRIMSON)
toc = [
    ("A", "Thông tin doanh nghiệp & câu chuyện thương hiệu"),
    ("B", "Bộ nhận diện thị giác — chốt design tokens (màu, font, bo góc, viền, bóng, hiệu ứng, họa tiết)"),
    ("C", "Thông tin liên hệ, cơ sở & vận hành"),
    ("D", "Cấu trúc website, menu & thứ tự trang"),
    ("E", "Chi tiết từng trang & từng section (thứ tự, nội dung, bố cục)"),
    ("F", "Nội dung dữ liệu chi tiết (dịch vụ, sản phẩm, đội ngũ, cơ sở, đánh giá, FAQ, bài viết)"),
    ("G", "Hình ảnh, video & media"),
    ("H", "Chuyển đổi: form đặt lịch, CTA, nút nổi"),
    ("I", "SEO, meta, domain & kỹ thuật hiển thị (không DB)"),
    ("J", "Giọng nội dung & copywriting"),
]
for letter, title in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, f"PHẦN {letter}.  ", bold=True, color=GOLD)
    run(p, title, size=10.5, color=INK)

doc.add_page_break()

# ============================================================
#  PHẦN A — DOANH NGHIỆP & THƯƠNG HIỆU
# ============================================================
part_heading("A", "Thông tin doanh nghiệp & câu chuyện thương hiệu")

sub_heading("A.1", "Định danh thương hiệu")
fill_table([
    ("Tên thương hiệu hiển thị", "Hiện đang dùng: “Y Viện Toplink”. Giữ nguyên hay đổi?"),
    ("Tên đầy đủ / tên công ty", "Dùng cho footer, hồ sơ pháp lý"),
    ("Tên viết tắt / cách gọi ngắn", "vd: Toplink"),
    ("Slogan / tagline chính", "Hiện: “Dưỡng thân từ gốc, phục hồi từ tâm”"),
    ("Slogan phụ (nếu có)", ""),
    ("Câu định vị 1 dòng", "Toplink là …? vd: Y Viện dưỡng thân Đông y cao cấp"),
    ("Lĩnh vực / ngành nghề", "Đông y, dưỡng sinh, trị liệu, spa sức khỏe…"),
    ("Năm thành lập / kinh nghiệm", "Chỉ ghi số thật"),
])

sub_heading("A.2", "Câu chuyện & giá trị")
open_block("Câu chuyện thương hiệu (About)", lines=4, hint="2–4 câu: vì sao Toplink ra đời, dành cho ai")
open_block("Sứ mệnh", lines=2)
open_block("Tầm nhìn", lines=2)
open_block("Giá trị cốt lõi (3–5 ý)", lines=3, hint="vd: Tận tâm · Bài bản · An toàn · Cá nhân hóa")
open_block("Cam kết với khách hàng", lines=2)

sub_heading("A.3", "Logo & tài sản nhận diện")
choice("Đã có logo chính thức chưa?", ["Đã có file vector (AI/SVG/PDF)", "Chỉ có file ảnh (PNG/JPG)", "Chưa có — cần thiết kế"])
fill_table([
    ("Đường dẫn / nơi gửi file logo", "Hiện web đang dùng chữ “Y” trong vòng triện làm tạm"),
    ("Logo có sẵn biến thể nào", "ngang / dọc / chỉ icon / nền sáng / nền tối"),
    ("Màu logo", "đỏ + vàng / vàng trên nền đỏ / khác"),
    ("Favicon (icon tab trình duyệt)", "đã có / dùng chữ Y / cần làm"),
    ("Có dùng con dấu / triện / hoa văn riêng?", ""),
])

# ============================================================
#  PHẦN B — BỘ NHẬN DIỆN THỊ GIÁC
# ============================================================
part_heading("B", "Bộ nhận diện thị giác — chốt design tokens")
note("Đây là các giá trị web ĐANG dùng. Hãy xác nhận “Giữ” hoặc ghi giá trị mới vào ô điền. "
     "Bạn không cần hiểu code — chỉ cần duyệt cảm giác màu/độ bo/độ động.", label="Quan trọng")

sub_heading("B.1", "Bảng màu (mã HEX đang dùng)")
note("Tỷ lệ khuyến nghị: nền kem/ngà 55–65% · đỏ 15–20% · nâu/đen 10–15% · vàng 5–8%. "
     "Vàng chỉ dùng làm điểm nhấn, không phủ mảng lớn.")
fill_table([
    ("Đỏ chủ đạo — #95131f", "Giữ / đổi sang HEX: ______"),
    ("Đỏ sâu (sơn mài) — #5c0a11", "Giữ / đổi: ______"),
    ("Đỏ nhấn (hotline) — #c70002", "Giữ / đổi: ______"),
    ("Vàng kim — #c8a02e", "Giữ / đổi: ______"),
    ("Vàng champagne (sáng) — #fdd79a", "Giữ / đổi: ______"),
    ("Nâu gỗ — #502c1e", "Giữ / đổi: ______"),
    ("Xanh ngọc (jade) — #2f5d50", "Giữ / bỏ / đổi: ______"),
    ("Kem nền — #ede7d3", "Giữ / đổi: ______"),
    ("Ngà nền trang — #fbf7f0", "Giữ / đổi: ______"),
    ("Mực chữ — #313131", "Giữ / đổi: ______"),
], col1="Vai trò màu (mã hiện tại)", col2="Quyết định")
choice("Tổng thể tông màu nên nghiêng về?",
       ["Giữ như hiện tại (đỏ sâu + vàng + gỗ + kem)",
        "Ấm/sáng hơn (nhiều kem, đỏ trầm hơn)",
        "Quyền uy hơn (nhiều đỏ sâu/đen sơn mài)",
        "Nhẹ nhàng hơn (giảm đỏ, tăng kem/nâu nhạt)"])

sub_heading("B.2", "Font chữ")
fill_table([
    ("Font tiêu đề (heading)", "Hiện: Playfair Display. Đề xuất bộ NĐ: Cormorant Garamond. Chọn: ______"),
    ("Font nội dung (body)", "Hiện: Be Vietnam Pro. Giữ?"),
    ("Font nhấn/calligraphy", "Hiện: Dancing Script (chữ viết tay). Giữ / đổi Cormorant Italic / bỏ"),
    ("Cỡ chữ body", "Hiện 18–20px (lớn cho khách trung niên). Giữ?"),
    ("Tiêu đề có IN HOA không", "menu/CTA hiện dùng in hoa giãn chữ"),
])

sub_heading("B.3", "Hình khối, bo góc, viền & đổ bóng")
fill_table([
    ("Bo góc nút bấm", "Hiện: 4px (gần vuông, không bo tròn pill). Giữ / đổi: ______"),
    ("Bo góc thẻ (card)", "Hiện: 8px. Giữ / đổi: ______"),
    ("Bo góc khối lớn (hero/CTA)", "Hiện: 12px. Giữ / đổi: ______"),
    ("Kiểu viền nhấn", "Hiện: viền vàng kim mảnh + khung triện. Giữ?"),
    ("Đổ bóng", "Hiện: bóng mềm nhẹ (không bóng nặng). Giữ / đậm hơn / phẳng hoàn toàn"),
])
choice("Cảm giác hình khối mong muốn?",
       ["Vuông vức, cổ điển, có khung triện (như hiện tại)",
        "Mềm hơn một chút (bo góc lớn hơn)",
        "Tối giản phẳng, ít viền trang trí"])

sub_heading("B.4", "Hiệu ứng động (animation)")
note("Web hiện chỉ dùng fade-up nhẹ khi cuộn, tôn trọng chế độ giảm chuyển động. Nhóm khách trung niên ưu tiên ít hiệu ứng.")
choice("Mức độ hiệu ứng động mong muốn?",
       ["Tối giản — chỉ hiện mượt khi cuộn (khuyến nghị, như hiện tại)",
        "Vừa — thêm hover, đếm số, ảnh trượt nhẹ",
        "Nhiều — parallax, hoa văn chuyển động, video nền"])
fill_table([
    ("Tốc độ chuyển động", "nhanh/dứt khoát hay chậm/êm?"),
    ("Có muốn hiệu ứng hover nổi bật trên thẻ dịch vụ?", ""),
    ("Có muốn số liệu đếm tăng (counter) ở khối niềm tin?", ""),
    ("Có dùng video nền ở hero không?", "nếu có: gửi video / để ảnh tĩnh"),
])

sub_heading("B.5", "Họa tiết & hoa văn Đông phương")
choice("Họa tiết muốn đưa vào web (nền section, viền, icon)?", multi=True, options=[
    "Mây lành", "Sóng nước", "Sen cách điệu", "Khung triện vuông",
    "Vân gỗ / giấy sắc thuốc", "Đường chỉ vàng", "Trống đồng / hoa văn Việt",
    "Mái ngói âm dương", "Bình hồ lô / dược bình", "Lá thuốc / rễ thuốc",
    "Không dùng họa tiết — giữ tối giản",
])
fill_table([
    ("Kiểu icon", "line-art mảnh vàng / icon đặc / icon hiện tại (lucide). Chọn: ______"),
    ("Mức độ họa tiết", "tiết chế (điểm nhấn) hay phủ nhiều?"),
])

# ============================================================
#  PHẦN C — LIÊN HỆ & CƠ SỞ
# ============================================================
part_heading("C", "Thông tin liên hệ, cơ sở & vận hành")

sub_heading("C.1", "Kênh liên hệ chính")
fill_table([
    ("Hotline chính", "Hiện: 0968 824 386 — đúng chưa?"),
    ("Hotline phụ (nếu có)", ""),
    ("Số/link Zalo", "Hiện đang trỏ tạm zalo.me"),
    ("Số WhatsApp/Viber (nếu có)", ""),
    ("Email liên hệ", ""),
    ("Fanpage Facebook", ""),
    ("Instagram / TikTok / YouTube", ""),
    ("Giờ làm việc", "Hiện: 08:00–21:00 hằng ngày"),
])

sub_heading("C.2", "Hệ thống cơ sở")
note("Liệt kê tất cả cơ sở thật. Web hiện có 2 cơ sở mẫu (Quận 1 & Quận 3).")
grid_table(["Tên cơ sở", "Địa chỉ đầy đủ", "SĐT", "Giờ mở cửa", "Link Google Maps"], 5,
           widths=[3.2, 5.0, 2.4, 2.4, 2.0])
choice("Có muốn hiển thị bản đồ Google Maps nhúng cho từng cơ sở?", ["Có", "Không", "Chỉ cơ sở chính"])

# ============================================================
#  PHẦN D — CẤU TRÚC & MENU
# ============================================================
part_heading("D", "Cấu trúc website, menu & thứ tự trang")

sub_heading("D.1", "Các trang hiện có & trang cần thêm")
note("Web hiện có: Trang chủ · Giới thiệu · Dịch vụ (+ chi tiết) · Quy trình trị liệu · Sản phẩm · Không gian · Tin tức · Đặt lịch · Liên hệ.")
choice("Trang muốn BỔ SUNG thêm?", multi=True, options=[
    "Học viện / Đào tạo KTV", "Nhượng quyền / Hợp tác", "Bảng giá tổng hợp",
    "Đội ngũ chuyên viên", "Tuyển dụng", "Chứng nhận & pháp lý",
    "Câu hỏi thường gặp (trang riêng)", "Ưu đãi / Khuyến mãi",
    "Thư viện ảnh (gallery)", "Không thêm — giữ nguyên",
])
choice("Trang muốn BỎ / gộp lại?", multi=True, options=[
    "Giữ tất cả", "Gộp Không gian vào Giới thiệu", "Gộp Tin tức vào trang khác",
    "Tạm ẩn Sản phẩm", "Tạm ẩn Tin tức",
])

sub_heading("D.2", "Thứ tự menu chính (header)")
note("Hiện tại header gồm các mục cố định. Hãy đánh số 1→N theo thứ tự mong muốn, hoặc ghi danh sách mới.")
open_block("Thứ tự menu chính mong muốn", lines=3,
           hint="vd: Trang chủ · Về Toplink · Liệu trình · Công nghệ · Đào tạo · Sản phẩm · Nhượng quyền · Đặt lịch")
fill_table([
    ("Nút nổi bật nhất trên header", "Hiện: “Đặt lịch”. Giữ?"),
    ("Có hiện hotline ngay trên header?", ""),
    ("Header có dính (sticky) khi cuộn?", "Hiện: có"),
    ("Menu có submenu/dropdown không?", "vd Dịch vụ xổ ra các nhóm"),
])

sub_heading("D.3", "Chân trang (footer)")
choice("Footer cần có khối nào?", multi=True, options=[
    "Logo + giới thiệu ngắn", "Menu nhanh", "Danh sách cơ sở + địa chỉ",
    "Hotline / Zalo / Email", "Mạng xã hội", "Giờ làm việc",
    "Bản đồ", "Đăng ký nhận tin", "Chứng nhận / pháp lý", "Bản quyền © Toplink",
])
open_block("Dòng giới thiệu ngắn ở footer", lines=2)

# ============================================================
#  PHẦN E — TỪNG TRANG & SECTION
# ============================================================
part_heading("E", "Chi tiết từng trang & từng section")
note("Phần quan trọng nhất. Với mỗi section: xác nhận GIỮ/BỎ/SỬA, thứ tự, và nội dung. "
     "Section đánh dấu (MỚI) là gợi ý bổ sung theo bộ nhận diện.")

# ---- E.1 TRANG CHỦ ----
sub_heading("E.1", "TRANG CHỦ — thứ tự & nội dung các section")
para("Thứ tự section hiện tại của trang chủ:", bold=True, size=10.5, color=WOOD)
order_now = [
    "Hero (khối đỏ + tiêu đề lớn + 2 nút)", "Bộ chọn nhu cầu (Need selector)",
    "Vì sao chọn Toplink (4 ô)", "Dịch vụ nổi bật (3 thẻ)",
    "Công nghệ & thiết bị (khối nâu gỗ)", "Quy trình trị liệu (4/8 bước)",
    "Không gian 4 tầng", "Cảm nhận khách hàng (reviews)",
    "Câu hỏi thường gặp (FAQ)", "Khối CTA cuối (đặt lịch/Zalo)",
]
for i, s in enumerate(order_now, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Cm(0.4)
    run(p, f"{i}. ", bold=True, color=GOLD); run(p, s, size=10, color=INK)
spacer(4)
choice("Section MỚI muốn thêm vào trang chủ?", multi=True, options=[
    "Khối số liệu/niềm tin (số khách, số KTV, năm KN)",
    "Khối Học viện / Đào tạo", "Khối Nhượng quyền / Hợp tác",
    "Khối Sản phẩm nổi bật", "Khối Tin tức/Kiến thức mới nhất",
    "Khối Đối tác / chứng nhận / báo chí", "Khối Bản đồ hệ thống cơ sở",
    "Khối Video giới thiệu", "Khối Ưu đãi đang chạy",
    "Không thêm",
])
open_block("Thứ tự section trang chủ mong muốn (đánh số lại nếu muốn đổi)", lines=3)

# Hero chi tiết
para("◆ Section HERO (đầu trang)", bold=True, size=11, color=CRIMSON)
fill_table([
    ("Tiêu đề lớn (headline)", "Hiện: “Dưỡng thân từ gốc, phục hồi từ tâm”"),
    ("Mô tả dưới tiêu đề", "Hiện: 2 dòng về Đông y + công nghệ"),
    ("Nhãn nhỏ phía trên (eyebrow)", "Hiện: “Y Viện Dưỡng Thân · Tỉnh Thức”"),
    ("Nút chính (CTA 1)", "Hiện: “Đặt lịch tư vấn”"),
    ("Nút phụ (CTA 2)", "Hiện: “Xem liệu trình”"),
    ("Ảnh/visual bên phải", "Hiện: ô giữ chỗ chờ ảnh thật. Gửi ảnh nào?"),
    ("Nền hero", "Hiện: khối đỏ sơn mài. Giữ / dùng ảnh nền / video"),
])

# Vì sao chọn
para("◆ Section “Vì sao chọn Toplink” (4 ô lý do)", bold=True, size=11, color=CRIMSON)
note("Hiện có 4 ô: Đông y dưỡng sinh · Cá nhân hóa · Công nghệ cao · An toàn minh bạch.")
grid_table(["Tiêu đề ô lý do", "Mô tả ngắn (1 dòng)"], 5, widths=[5.0, 10.0])

# Khối niềm tin (mới)
para("◆ Section số liệu / niềm tin (MỚI — nếu chọn ở trên)", bold=True, size=11, color=CRIMSON)
note("Chỉ dùng số THẬT. Chưa có thì dùng cam kết chất lượng thay cho con số.")
grid_table(["Con số / cam kết", "Nhãn mô tả"], 4, widths=[5.0, 10.0])

# ---- E.2 DỊCH VỤ ----
sub_heading("E.2", "TRANG DỊCH VỤ & trang chi tiết dịch vụ")
fill_table([
    ("Cách nhóm dịch vụ", "Hiện: Cơ bản / Nâng cao / Chuyên sâu. Giữ / đổi theo vùng cơ thể?"),
    ("Có hiển thị giá không?", "Hiện: hiện giá “từ …đ”. Giữ / ẩn / ghi “liên hệ”"),
    ("Có bộ lọc theo nhu cầu?", "Hiện: có bộ chọn nhu cầu"),
    ("Mỗi thẻ dịch vụ hiện gì", "Hiện: tên, mô tả, thời lượng, giá, nút. Thêm/bớt?"),
])
choice("Trang CHI TIẾT mỗi dịch vụ cần khối nào?", multi=True, options=[
    "Ảnh liệu trình", "Mô tả dài", "Thời lượng & giá", "Phù hợp với ai",
    "Lưu ý / chống chỉ định", "Các bước thực hiện", "Cảm nhận sau liệu trình",
    "Dịch vụ liên quan", "Đánh giá khách", "Nút đặt lịch/Zalo", "FAQ riêng",
])

# ---- E.3 GIỚI THIỆU ----
sub_heading("E.3", "TRANG GIỚI THIỆU (Về Toplink)")
choice("Khối nội dung cho trang Giới thiệu?", multi=True, options=[
    "Câu chuyện thương hiệu", "Sứ mệnh / tầm nhìn / giá trị", "Đội ngũ chuyên viên",
    "Triết lý Đông y dưỡng sinh", "Không gian Y Viện", "Chứng nhận / giấy phép",
    "Dòng thời gian phát triển", "Cam kết với khách hàng", "Hình ảnh thực tế",
])

# ---- E.4 KHÔNG GIAN ----
sub_heading("E.4", "TRANG KHÔNG GIAN")
note("Hiện mô tả 4 tầng: Tĩnh · Thông · Dưỡng · Tỉnh.")
fill_table([
    ("Mô hình không gian thực tế", "Có đúng 4 tầng không? Mô tả từng khu"),
    ("Ảnh không gian", "Gửi ảnh từng khu/phòng"),
    ("Tên các phòng trị liệu", "vd: Phòng Dưỡng Tâm, Khai Thông, An Nhiên…"),
])

# ---- E.5 SẢN PHẨM ----
sub_heading("E.5", "TRANG SẢN PHẨM")
choice("Trang sản phẩm hoạt động kiểu gì?", options=[
    "Chỉ giới thiệu + tư vấn qua Zalo (không bán online)",
    "Có giá + nút đặt mua qua Zalo",
    "Bán online đầy đủ (giỏ hàng) — ngoài phạm vi demo hiện tại",
])
fill_table([
    ("Nhóm sản phẩm", "Hiện: Thảo dược / Hỗ trợ tại nhà / Máy sức khỏe"),
    ("Mỗi sản phẩm hiện gì", "tên, nhóm, giá, mô tả, công dụng, cách dùng, lưu ý"),
])

# ---- E.6 TIN TỨC ----
sub_heading("E.6", "TRANG TIN TỨC / KIẾN THỨC")
fill_table([
    ("Chuyên mục bài viết", "Hiện: Kiến thức sức khỏe / Đông y dưỡng sinh / Quy trình"),
    ("Tần suất đăng bài", ""),
    ("Ai viết nội dung", "Toplink cung cấp / cần hỗ trợ viết"),
])

# ---- E.7 ĐẶT LỊCH ----
sub_heading("E.7", "TRANG ĐẶT LỊCH")
para("→ Chi tiết trường form ở Phần H.", italic=True, size=9.5, color=INK_SOFT)

# ---- E.8 ĐÀO TẠO / NHƯỢNG QUYỀN (mới) ----
sub_heading("E.8", "TRANG ĐÀO TẠO & NHƯỢNG QUYỀN (MỚI — nếu chọn ở Phần D)")
open_block("Thông điệp trang Đào tạo", lines=2, hint="vd: Đào tạo KTV trị liệu Đông y bài bản")
open_block("Quyền lợi / lộ trình học viên", lines=3)
open_block("Thông điệp trang Nhượng quyền", lines=2)
open_block("Hỗ trợ dành cho đối tác", lines=3, hint="vận hành, đào tạo, set-up, marketing, sản phẩm")

# ============================================================
#  PHẦN F — NỘI DUNG DỮ LIỆU
# ============================================================
part_heading("F", "Nội dung dữ liệu chi tiết")
note("Điền dữ liệu THẬT để thay nội dung mẫu. Thêm dòng nếu thiếu.")

sub_heading("F.1", "Danh sách dịch vụ / liệu trình đầy đủ")
grid_table(["Tên dịch vụ", "Nhóm", "Thời lượng", "Giá từ", "Phù hợp với ai", "Lưu ý / chống chỉ định"],
           8, widths=[3.0, 1.8, 1.6, 1.6, 3.5, 3.5])
note("Mỗi dịch vụ còn cần: mô tả ngắn 1–2 dòng, các bước thực hiện, cảm nhận sau buổi. "
     "Có thể gửi riêng theo mẫu này.")

sub_heading("F.2", "Danh sách sản phẩm")
grid_table(["Tên sản phẩm", "Nhóm", "Giá từ", "Công dụng chính", "Cách dùng / lưu ý"],
           6, widths=[3.2, 2.0, 1.8, 4.0, 4.0])

sub_heading("F.3", "Đội ngũ chuyên viên / kỹ thuật viên")
grid_table(["Họ tên / chức danh", "Chuyên môn", "Kinh nghiệm / chứng chỉ"], 5,
           widths=[4.5, 5.0, 5.5])

sub_heading("F.4", "Đánh giá khách hàng thật")
note("Chỉ dùng feedback thật, có thể ẩn bớt tên. Web hiện có 3 review mẫu.")
grid_table(["Tên (có thể viết tắt)", "Vai trò / nghề", "Nội dung cảm nhận"], 5,
           widths=[3.5, 3.5, 8.0])

sub_heading("F.5", "Câu hỏi thường gặp (FAQ)")
grid_table(["Câu hỏi", "Câu trả lời"], 6, widths=[5.5, 9.5])

sub_heading("F.6", "Bài viết / kiến thức dự kiến")
grid_table(["Tiêu đề bài", "Chuyên mục", "Tóm tắt 1 dòng"], 5, widths=[5.0, 3.5, 6.5])

# ============================================================
#  PHẦN G — HÌNH ẢNH & MEDIA
# ============================================================
part_heading("G", "Hình ảnh, video & media")
note("Ảnh thật là yếu tố quyết định cảm giác cao cấp. Tránh ảnh stock bác sĩ Tây, ảnh quá spa nữ tính, ảnh trắng-xanh bệnh viện.")
fill_table([
    ("Ảnh Hero trang chủ", "1–2 ảnh không gian/đội ngũ chất lượng cao"),
    ("Ảnh không gian từng khu/phòng", ""),
    ("Ảnh đội ngũ / KTV (đồng phục)", ""),
    ("Ảnh từng dịch vụ / liệu trình", ""),
    ("Ảnh sản phẩm (nền sạch, ánh ấm)", ""),
    ("Ảnh khách trải nghiệm (có cho phép)", ""),
    ("Ảnh chứng nhận / giấy phép / đào tạo", ""),
    ("Video giới thiệu (nếu có)", "link / file"),
    ("Logo đối tác / báo chí (nếu có)", ""),
])
choice("Toplink sẽ cung cấp ảnh thế nào?", options=[
    "Có sẵn bộ ảnh chuyên nghiệp — sẽ gửi",
    "Có ảnh nhưng cần chỉnh sửa", "Cần Toplink chụp mới / hỗ trợ",
    "Tạm dùng ảnh minh họa rồi thay sau",
])

# ============================================================
#  PHẦN H — CHUYỂN ĐỔI
# ============================================================
part_heading("H", "Chuyển đổi: form đặt lịch, CTA, nút nổi")

sub_heading("H.1", "Form đặt lịch")
note("Khách ngành trị liệu thích được tư vấn nhanh — form nên NGẮN. Web hiện có stepper nhiều bước.")
choice("Các trường cần có trong form đặt lịch?", multi=True, options=[
    "Họ tên", "Số điện thoại", "Nhu cầu / vấn đề cơ thể", "Chọn dịch vụ",
    "Chọn cơ sở", "Ngày mong muốn", "Khung giờ mong muốn", "Ghi chú thêm",
    "Email (không bắt buộc)",
])
fill_table([
    ("Trường BẮT BUỘC tối thiểu", "vd: chỉ Tên + SĐT"),
    ("Form gửi đi đâu", "Hiện chưa nối DB. Tạm: hiện thông báo / chuyển sang Zalo / gửi email?"),
    ("Tin nhắn xác nhận sau khi gửi", "vd: “Toplink sẽ gọi lại trong 15 phút”"),
])

sub_heading("H.2", "Nút kêu gọi hành động (CTA) & nút nổi")
choice("CTA nên xuất hiện ở những vị trí nào?", multi=True, options=[
    "Hero", "Sau danh sách dịch vụ", "Sau quy trình", "Sau đánh giá khách",
    "Cuối trang (khối CTA)", "Thanh nổi cố định trên mobile", "Header",
])
fill_table([
    ("Microcopy CTA chính", "Hiện: “Đặt lịch tư vấn”. Giữ / đổi: ______"),
    ("Thanh nổi mobile gồm nút nào", "Hiện: Gọi · Zalo · Đặt lịch · Chỉ đường"),
])

# ============================================================
#  PHẦN I — SEO / META / KỸ THUẬT
# ============================================================
part_heading("I", "SEO, meta, domain & hiển thị (không DB)")
fill_table([
    ("Tên miền (domain)", "đã có / cần mua. vd: toplink.vn"),
    ("Tiêu đề trang (Title) mặc định", "hiển thị trên tab & Google"),
    ("Mô tả (Meta description)", "1–2 câu mô tả cho Google"),
    ("Từ khóa SEO mục tiêu", "vd: trị liệu Đông y, dưỡng sinh, cổ vai gáy…"),
    ("Ảnh chia sẻ mạng XH (OG image)", "ảnh hiện khi gửi link lên FB/Zalo"),
    ("Ngôn ngữ", "Tiếng Việt / song ngữ?"),
    ("Google Analytics / Pixel", "có cài theo dõi không?"),
    ("Google Maps / Google Business", "link địa điểm"),
    ("Chứng chỉ bảo mật (HTTPS)", "thường mặc định có"),
])

# ============================================================
#  PHẦN J — GIỌNG NỘI DUNG
# ============================================================
part_heading("J", "Giọng nội dung & copywriting")
note("Giọng đề xuất: chuyên môn nhưng không khô · Đông y nhưng không mê tín · cao cấp nhưng không xa cách · "
     "chuyển đổi mạnh nhưng không bán hàng lộ liễu.")
choice("Cách xưng hô với khách trên web?", options=[
    "“chị/anh” (ấm, gần gũi — như hiện tại)", "“quý khách” (trang trọng)",
    "“bạn” (trẻ trung)", "Kết hợp tùy ngữ cảnh",
])
choice("Tông giọng tổng thể?", multi=True, options=[
    "Ấm áp, chăm sóc", "Trang trọng, cao cấp", "Tỉnh thức, an nhiên",
    "Chuyên môn, đáng tin", "Truyền cảm hứng",
])
open_block("Từ khóa / cụm từ THƯƠNG HIỆU muốn dùng", lines=2,
           hint="vd: Dưỡng thân · Khai thông · Cân bằng · Phục hồi · Thân–Tâm–Trí")
open_block("Từ ngữ CẦN TRÁNH", lines=2,
           hint="vd: chữa khỏi 100%, cam kết khỏi bệnh, thần dược, trị dứt điểm")
open_block("Câu slogan / triết lý muốn xuất hiện nổi bật", lines=2)

# ---- Khối kết ----
doc.add_page_break()
pend = doc.add_paragraph(); pend.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(pend, "— HẾT BIỂU MẪU —", bold=True, size=12, color=GOLD, font=HEAD_FONT)
pend.paragraph_format.space_after = Pt(4)
pe2 = doc.add_paragraph(); pe2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(pe2, "Sau khi điền xong, gửi lại file này kèm bộ ảnh & logo (nếu có). "
         "Đội thiết kế sẽ dựng bản nâng cấp theo đúng lựa chọn của bạn.",
    italic=True, size=10, color=INK_SOFT)

# ---------- Lưu ----------
out = r"f:\Codex\Yvien Hotlink Website\Z-NeededUpdate\FORM-Hoan-Thien-Website-Y-Vien-Toplink.docx"
doc.save(out)
print("SAVED:", out)
